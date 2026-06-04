from __future__ import annotations

import io
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.tenant import scoped
from app.models.document_registry import DocumentRegistry
from app.models.organization import Organization
from app.models.service_catalog import ServiceCatalog
from app.rag.chunker import DocumentChunk, chunk_text
from app.rag.constants import RAG_NAMESPACES
from app.rag.indexer import KnowledgeIndexer
from app.schemas.service_catalog import (
    ServiceCatalogCreate,
    ServiceCatalogItem,
    ServiceCatalogUpdate,
)
from app.services.service_catalog_service import (
    ServiceCatalogService,
    build_service_rag_text,
)

_PRICING_FAQ_ID = "pricing::GENERAL_FAQ"
_PRICING_FAQ_TEXT = (
    "Pricing FAQ: All prices listed are estimates based on typical jobs. "
    "Exact quotes require an on-site technician visit because conditions, "
    "parts, and equipment age vary. Diagnostic fees may be waived when you "
    "proceed with a repair — ask your technician. Emergency surcharges apply "
    "for after-hours and same-day calls."
)


def _slugify(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", value.lower()).strip("-")
    return slug or "document"


_PAGE_BREAK = "\n--- PAGE BREAK ---\n"


def _extract_text(filename: str, content: bytes) -> tuple[str, list[str]]:
    """Return (text, warnings) for supported document types."""
    lower = filename.lower()
    warnings: list[str] = []

    if lower.endswith(".csv"):
        raise ValueError(
            "CSV files cannot be uploaded as knowledge documents. "
            "Use POST /api/v1/imports/{org_id}/customers or /equipment instead."
        )
    if lower.endswith((".md", ".txt", ".text")):
        return content.decode("utf-8", errors="replace"), warnings
    if lower.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = _PAGE_BREAK.join(p.strip() for p in pages if p.strip())
        if not text.strip():
            warnings.append(
                "No extractable text found in PDF (it may be scanned/image-only)."
            )
        return text, warnings
    if lower.endswith((".docx", ".doc")):
        from docx import Document

        doc = Document(io.BytesIO(content))
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            line = paragraph.text.strip()
            if not line:
                continue
            style = (paragraph.style.name or "").lower()
            if "heading" in style:
                parts.append(f"## {line}")
            else:
                parts.append(line)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n\n".join(parts)
        if not text.strip():
            warnings.append("No extractable text found in Word document.")
        return text, warnings

    raise ValueError(
        f"Unsupported file type for '{filename}'. "
        "Supported: .pdf, .docx, .md, .txt"
    )


class KnowledgeService:
    def __init__(self, db: AsyncSession, indexer: KnowledgeIndexer | None = None) -> None:
        self.db = db
        self.indexer = indexer or KnowledgeIndexer()
        self.catalog_service = ServiceCatalogService(db)

    async def _get_org(self, org_id: uuid.UUID) -> Organization | None:
        return await self.db.get(Organization, org_id)

    async def list_namespaces(self, org_id: uuid.UUID) -> list[dict[str, Any]]:
        org = await self._get_org(org_id)
        if org is None:
            return []
        counts = self.indexer.get_namespace_counts()
        return [
            {"namespace": ns, "vector_count": counts.get(ns, 0)}
            for ns in sorted(RAG_NAMESPACES)
        ]

    async def list_documents(self, org_id: uuid.UUID) -> list[dict[str, Any]]:
        stmt = scoped(
            select(DocumentRegistry)
            .where(DocumentRegistry.is_active.is_(True))
            .order_by(DocumentRegistry.uploaded_at.desc()),
            DocumentRegistry,
            org_id,
        )
        rows = (await self.db.execute(stmt)).scalars().all()
        return [
            {
                "doc_id": str(row.doc_id),
                "document_id": row.document_id,
                "filename": row.filename,
                "namespace": row.namespace,
                "chunk_count": row.chunk_count,
                "file_size_bytes": row.file_size_bytes,
                "mime_type": row.mime_type,
                "uploaded_at": row.uploaded_at.isoformat(),
                "last_indexed_at": row.last_indexed_at.isoformat(),
            }
            for row in rows
        ]

    async def upload_document(
        self,
        org_id: uuid.UUID,
        filename: str,
        content: bytes,
        namespace: str = "faq_general",
        document_id: str | None = None,
        mime_type: str | None = None,
        chunking_strategy: str = "paragraph",
    ) -> dict[str, Any]:
        org = await self._get_org(org_id)
        if org is None:
            raise ValueError(f"Organization {org_id} not found")
        if namespace not in RAG_NAMESPACES:
            raise ValueError(f"Invalid namespace: {namespace}")

        doc_id = document_id or _slugify(Path(filename).stem)
        text, extract_warnings = _extract_text(filename, content)
        chunks = chunk_text(
            text=text,
            source=filename,
            namespace=namespace,
            strategy=chunking_strategy,
        )
        if not chunks:
            chunks = [
                DocumentChunk(
                    text=text[:512] or filename,
                    source=filename,
                    namespace=namespace,
                    chunk_index=0,
                )
            ]

        chunk_ids = [f"{doc_id}::{index}" for index in range(len(chunks))]
        await self.indexer.delete_by_id_prefix(namespace, f"{doc_id}::")
        chunks_indexed = await self.indexer.index_chunks(chunks, chunk_ids=chunk_ids)

        now = datetime.now(timezone.utc)
        existing = (
            await self.db.execute(
                select(DocumentRegistry).where(
                    DocumentRegistry.org_id == org_id,
                    DocumentRegistry.document_id == doc_id,
                )
            )
        ).scalar_one_or_none()

        if existing:
            existing.filename = filename
            existing.namespace = namespace
            existing.chunk_count = chunks_indexed
            existing.file_size_bytes = len(content)
            existing.mime_type = mime_type
            existing.last_indexed_at = now
            existing.is_active = True
        else:
            self.db.add(
                DocumentRegistry(
                    org_id=org_id,
                    document_id=doc_id,
                    filename=filename,
                    namespace=namespace,
                    chunk_count=chunks_indexed,
                    file_size_bytes=len(content),
                    mime_type=mime_type,
                    uploaded_at=now,
                    last_indexed_at=now,
                    is_active=True,
                )
            )
        await self.db.flush()
        return {
            "document_id": doc_id,
            "chunks_indexed": chunks_indexed,
            "namespace": namespace,
            "warnings": extract_warnings,
        }

    async def delete_document(
        self, org_id: uuid.UUID, document_id: str
    ) -> dict[str, Any] | None:
        row = (
            await self.db.execute(
                select(DocumentRegistry).where(
                    DocumentRegistry.org_id == org_id,
                    DocumentRegistry.document_id == document_id,
                )
            )
        ).scalar_one_or_none()
        if row is None:
            return None
        await self.indexer.delete_by_id_prefix(row.namespace, f"{document_id}::")
        row.is_active = False
        await self.db.flush()
        return {"deleted": True, "document_id": document_id}

    async def index_service_to_pricing(self, row: ServiceCatalog) -> int:
        prefix = f"pricing::{row.service_code}"
        await self.indexer.delete_by_id_prefix("pricing", prefix)
        text = build_service_rag_text(row)
        chunk = DocumentChunk(
            text=text,
            source=f"service_catalog:{row.service_code}",
            namespace="pricing",
            chunk_index=0,
        )
        return await self.indexer.index_chunks(
            [chunk], chunk_ids=[f"{prefix}::0"]
        )

    async def index_pricing_faq(self) -> int:
        await self.indexer.delete_by_id_prefix("pricing", _PRICING_FAQ_ID)
        chunk = DocumentChunk(
            text=_PRICING_FAQ_TEXT,
            source="pricing_faq",
            namespace="pricing",
            chunk_index=0,
        )
        return await self.indexer.index_chunks([chunk], chunk_ids=[f"{_PRICING_FAQ_ID}::0"])

    async def create_service_catalog_entry(
        self, org_id: uuid.UUID, data: ServiceCatalogCreate
    ) -> ServiceCatalogItem:
        item = await self.catalog_service.create(org_id, data)
        row = await self.catalog_service.get_by_id(org_id, uuid.UUID(item.service_id))
        if row is not None and row.is_active:
            await self.index_service_to_pricing(row)
        return item

    async def update_service_catalog_entry(
        self,
        org_id: uuid.UUID,
        service_id: uuid.UUID,
        data: ServiceCatalogUpdate,
    ) -> ServiceCatalogItem | None:
        item = await self.catalog_service.update(org_id, service_id, data)
        if item is None:
            return None
        row = await self.catalog_service.get_by_id(org_id, service_id)
        if row is not None and row.is_active:
            await self.index_service_to_pricing(row)
        elif row is not None:
            await self.indexer.delete_by_id_prefix(
                "pricing", f"pricing::{row.service_code}"
            )
        return item

    async def list_service_catalog(self, org_id: uuid.UUID) -> list[ServiceCatalogItem]:
        return await self.catalog_service.list_all(org_id)
